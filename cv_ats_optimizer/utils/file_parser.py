"""Utilities for parsing uploaded files into clean text."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Callable

import docx
import pdfplumber
from PyPDF2 import PdfReader

from cv_ats_optimizer.utils.text_processor import clean_text


class FileParsingError(RuntimeError):
    """Raised when file parsing fails."""


def parse_txt(file_bytes: bytes) -> str:
    """Parse plain text files, trying utf-8 then latin-1 encodings."""

    for encoding in ("utf-8", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            cleaned = clean_text(text)
            if cleaned:
                return cleaned
        except UnicodeDecodeError:
            continue
    raise FileParsingError("Unable to decode text file. Please ensure UTF-8 or Latin-1 encoding.")


def parse_pdf(file_bytes: bytes) -> str:
    """Parse PDF files using pdfplumber with PyPDF2 fallback."""

    text_content = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text_content.append(page_text)
    if text_content:
        cleaned = clean_text("\n".join(text_content))
        if cleaned:
            return cleaned

    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    cleaned = clean_text("\n".join(pages_text))
    if not cleaned:
        raise FileParsingError("PDF appears to be empty or contains unsupported text layers.")
    return cleaned


def parse_docx(file_bytes: bytes) -> str:
    """Parse DOCX documents by iterating paragraphs and tables."""

    document = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    cleaned = clean_text("\n".join(parts))
    if not cleaned:
        raise FileParsingError("DOCX file contains no extractable text.")
    return cleaned


def parse_file(filename: str, file_bytes: bytes) -> str:
    """Route parsing based on file extension."""

    extension = Path(filename).suffix.lower()
    parsers: dict[str, Callable[[bytes], str]] = {
        ".txt": parse_txt,
        ".pdf": parse_pdf,
        ".docx": parse_docx,
    }
    if extension not in parsers:
        raise FileParsingError(f"Unsupported file extension: {extension}.")
    return parsers[extension](file_bytes)


__all__ = [
    "FileParsingError",
    "parse_txt",
    "parse_pdf",
    "parse_docx",
    "parse_file",
]
